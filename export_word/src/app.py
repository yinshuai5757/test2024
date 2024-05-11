import datetime
import shutil
import os
import re
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips
from flask import Flask, request, send_file
app = Flask(__name__)

template_file_path = "./template.docx"
#logging.basicConfig(level=logging.INFO,format="%(asctime)s - %(levelname)s:%(name)s - %(message)s",filename="/logs/word_logs.log")
# 接頭辞のリストを定義
prefixes = [
    ["第１　", "第２　", "第３　", "第４　", "第５　", "第６　", "第７　", "第８　", "第９　", "第１０　"],
    ["１　", "２　", "３　", "４　", "５　", "６　", "７　", "８　", "９　", "１０　"],
    ["⑴　", "⑵　", "⑶　", "⑷　", "⑸　", "⑹　", "⑺　", "⑻　", "⑼　", "⑽　"],
    ["ア　", "イ　", "ウ　", "エ　", "オ　", "カ　", "キ　", "ク　", "ケ　", "コ　"],
    ["（ア）　", "（イ）　", "（ウ）　", "（エ）　", "（オ）　", "（カ）　", "（キ）　", "（ク）　", "（ケ）　", "（コ）　"],
    ["a　", "b　", "c　", "d　", "e　", "f　", "g　", "h　", "i　", "j　"],
    ["（a）　", "（b）　", "（c）　", "（d）　", "（e）　", "（f）　", "（g）　", "（h）　", "（i）　", "（j）　"],
    ["Ⅰ　", "Ⅱ　", "Ⅲ　", "Ⅳ　", "Ⅴ　", "Ⅵ　", "Ⅶ　", "Ⅷ　", "Ⅸ　", "Ⅹ　"],
    ["α　", "β　", "γ　", "δ　", "ε　", "ζ　", "η　", "θ　", "ι　", "κ　"]
]

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        nodes_text = request.data.decode('utf-8')
        now = datetime.datetime.now()
        formatted_date = now.strftime("%Y%m%d%H%M%S")
        file_name = f"sample_{formatted_date}.docx"
        # カレントディレクトリを取得してファイル名と結合
        file_path = os.path.join(os.getcwd(), file_name)
        shutil.copy(template_file_path, file_path)

        export_to_word(nodes_text, file_path)

        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return str(e), 500

def export_to_word(nodes_text, file_path):
    document = Document(file_path)
    level_counters = [0, 0, 0, 0, 0, 0, 0, 0, 0]
    last_level = 1  # 最後に処理されたレベル

    for line in nodes_text.strip().split('\n'):
        # 行頭が "#" または "//" の場合、スキップする
        if line.startswith("#") or line.startswith("//"):
            continue
        # hidden title（行頭に "\**-"（正規表現）がある場合、スキップする）
        if re.match(r'^\**-', line):
            continue
        level = line.count('*')   # "*"の数でレベルを決定
        if level == 0:            # "*"がない場合（「との判決を求める。」のようなケース）、直近のレベルを使用し、level_counters は変更しない
            level = last_level
        elif last_level == level: # 兄弟ノードに遷移したら、ノードの level_counters をインクリメント
            level_counters[level] += 1
        elif last_level < level:  # 子ノードに遷移したら、子ノードの level_counters をインクリメント
            level_counters[level] += 1
        elif level < last_level:  # 親ノードに遷移したら、子ノードの level_counters をクリアし、ノードの level_counters をインクリメント
            for l in range(level + 1, len(level_counters)):
                level_counters[l] = 0
            level_counters[level] += 1
        else:                     # 存在しないケース
            level = last_level
        last_level = level
        write_node_to_doc(document, line, level, level_counters[level]-1) # counter の定義域は [0,n]

    document.save(file_path)

def write_node_to_doc(document, text, level, counter):
    node_text = ""
    layout = ""
    # title（章題）か bosy（本文）
    is_title = True
    # テキストに接頭辞を追加
    prefix = prefixes[level-1][counter] if level > 0 else ""
    if 0 < text.count('*'):
        text = text.replace('*', '').strip()
        if 0 < len(text):
            node_text = prefix + text
        is_title = True
    else: # '*' がなければ接頭辞は付与しないで、行末の "@str"（位置指定修飾）の有無を確認
        is_title = False
        layout = text.split("@")[-1]   # 行末の "@str"（位置指定修飾）を取得する
        layout = layout.rstrip("\r\n") # 行末の "\r" と "\n" を削除
        text = text[:text.rfind("@")]  # 行末の "@str" を削除する
        node_text = text

    # テキストを追加
    paragraph = document.add_paragraph()
    run = paragraph.add_run(node_text)
    font = run.font
    font.size = Pt(12)

    set_indentation(paragraph, level, is_title, node_text)
    if layout.isdigit():
        level = int(layout)
        set_indentation(paragraph, level, is_title, node_text) # 改めて indent を設定
    elif "right" in layout:
        # この行を右揃えにする
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif "break" in layout:
        # node_text の行末に改ページを挿入
        document.add_page_break()

def set_indentation(paragraph, level, is_title, node_text):
    # 240 twips 単位でのインデント量を計算
    left_indent_value = Twips(240 * level)
    if is_title:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = left_indent_value # 段落全体の indent をまず設定する
        paragraph_format.first_line_indent = -Twips(240) # left_indent からの差分を設定する
    else:
        if level < len(prefixes) and any(node_text.startswith(prefix) for prefix in prefixes[level]):
            # node_text の行頭に prefixes のいずれかがある場合の処理（ユーザが自分で章番号を記載した場合）
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = left_indent_value+Twips(240) # 段落全体の indent をまず設定する
            paragraph_format.first_line_indent = -Twips(240) # left_indent からの差分を設定する
        else:
            # 通常の場合
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = left_indent_value # 段落全体の indent をまず設定する
            paragraph_format.first_line_indent = Twips(240)  # left_indent からの差分を設定する

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=7008)
